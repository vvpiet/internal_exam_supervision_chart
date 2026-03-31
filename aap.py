import streamlit as st
import pandas as pd
from datetime import date, timedelta
from docx import Document
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import io
import os

st.set_page_config(page_title="Exam Supervision Chart", layout="wide")

# Custom CSS for styling
st.markdown("""
    <style>
    body {
        background-color: #f5f5f5;
    }
    .header-container {
        display: flex;
        align-items: center;
        justify-content: center;
        gap: 20px;
        padding: 20px;
        background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
        border-radius: 10px;
        margin-bottom: 30px;
        color: white;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    .header-text h1 {
        margin: 0;
        font-size: 32px;
        font-weight: bold;
        text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.3);
    }
    .footer-container {
        margin-top: 50px;
        padding: 20px;
        background: linear-gradient(135deg, #2a5298 0%, #1e3c72 100%);
        border-radius: 10px;
        color: white;
        text-align: center;
        font-size: 14px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    .footer-text {
        margin: 10px 0;
        font-weight: 500;
    }
    </style>
""", unsafe_allow_html=True)

# Header with Logo and Title
col1, col2 = st.columns([1, 4])
with col1:
    # Try to load logo if it exists, otherwise use a placeholder icon
    logo_path = "vvp_logo.png"
    if os.path.exists(logo_path):
        st.image(logo_path, width=100)
    else:
        # Using a professional university icon as placeholder
        st.image("https://img.icons8.com/color/200/000000/organization.png", width=100)
        st.caption("(Add vvp_logo.png to project folder)")
with col2:
    st.markdown("""
        <div style='color: #1e3c72; padding: 20px;'>
            <h1>VVP Institute of Engineering and Technology, Solapur</h1>
            <h3 style='color: #2a5298; margin-top: -10px;'>Internal Exam Supervision Chart Generator</h3>
        </div>
    """, unsafe_allow_html=True)

st.markdown("---")

# Sidebar Inputs
st.sidebar.header("Configuration")
st.sidebar.info("💡 To use your institute logo, save it as 'vvp_logo.png' in the project folder")

# Date Range Input
col1, col2 = st.sidebar.columns(2)
with col1:
    exam_start_date = st.date_input("Exam Start Date")
with col2:
    exam_end_date = st.date_input("Exam End Date")

# Holidays Input
st.sidebar.subheader("Holidays to Exclude")
holidays_input = st.sidebar.text_area(
    "Enter holidays (one date per line, format: DD-MM-YYYY)",
    "25-12-2024\n26-12-2024"
)

holidays = []
if holidays_input.strip():
    for holiday in holidays_input.strip().split("\n"):
        try:
            holiday_date = pd.to_datetime(holiday.strip(), format="%d-%m-%Y").date()
            holidays.append(holiday_date)
        except:
            st.sidebar.warning(f"Invalid date format: {holiday.strip()}")

st.sidebar.info(f"Excluded holidays: {len(holidays)} dates")

# Blocks per Day Configuration
st.sidebar.subheader("Blocks Per Day")
col1, col2 = st.sidebar.columns(2)
with col1:
    morning_blocks = st.number_input("Morning Blocks", min_value=1, value=2, step=1)
with col2:
    evening_blocks = st.number_input("Evening Blocks", min_value=0, value=2, step=1)

# Morning Time Slots
st.sidebar.write("**Morning Time Slots (comma separated)**")
morning_slots_input = st.sidebar.text_input(
    "Morning Slots",
    "11:00-12:00,12:15-1:15"
)

# Evening Time Slots
st.sidebar.write("**Evening Time Slots (comma separated)**")
evening_slots_input = st.sidebar.text_input(
    "Evening Slots",
    "2:00-3:00,3:15-4:15"
)

morning_time_slots = [s.strip() for s in morning_slots_input.split(",") if s.strip()]
evening_time_slots = [s.strip() for s in evening_slots_input.split(",") if s.strip()]

# Combine time slots based on block configuration
time_slots = []
if morning_blocks > 0:
    time_slots.extend(morning_time_slots[:morning_blocks])
if evening_blocks > 0:
    time_slots.extend(evening_time_slots[:evening_blocks])

# Faculty List Import from Excel
st.sidebar.subheader("Faculty List")
excel_file = st.sidebar.file_uploader("Upload Faculty List (Excel)", type=["xlsx", "xls"])

faculty_list = []

if excel_file is not None:
    try:
        df_faculty = pd.read_excel(excel_file)
        # Assume faculty names are in the first column
        faculty_list = [f.strip() for f in df_faculty.iloc[:, 0].astype(str) if f.strip() and f.strip() != 'nan']
        st.sidebar.success(f"Loaded {len(faculty_list)} faculty members from Excel")
    except Exception as e:
        st.sidebar.error(f"Error reading Excel file: {e}")
else:
    st.sidebar.info("Please upload an Excel file with faculty names in the first column")
    
    # Fallback: Manual entry
    st.sidebar.write("**Or enter manually (comma separated)**")
    senior_faculty = st.sidebar.text_area(
        "Senior Faculty",
        "Dr. Patil,Dr. Mehta"
    ).split(",")

    junior_faculty = st.sidebar.text_area(
        "Junior Faculty",
        "Prof. Shah,Prof. Kumar,Prof. Rao"
    ).split(",")

    faculty_list = [f.strip() for f in senior_faculty] + [f.strip() for f in junior_faculty]
    faculty_list = [f for f in faculty_list if f]


if st.button("Generate Chart"):
    if not faculty_list:
        st.error("Please provide faculty list (either upload Excel or enter manually)")
    elif exam_start_date > exam_end_date:
        st.error("Exam start date must be before end date")
    elif len(time_slots) == 0:
        st.error("Please configure at least one time slot")
    else:
        # Generate date range excluding holidays
        current_date = exam_start_date
        date_range = []
        while current_date <= exam_end_date:
            if current_date not in holidays:
                date_range.append(current_date)
            current_date += timedelta(days=1)

        total_slots = len(date_range) * len(time_slots)
        num_faculty = len(faculty_list)
        
        # Create assignment mapping: rotating through faculty
        faculty_assignments = {}  # {faculty_name: [(date1, slot_idx1), (date2, slot_idx2), ...]}
        faculty_index = 0
        
        # Assign faculty by rotating through all available slots
        for exam_date in date_range:
            for slot_idx in range(len(time_slots)):
                assigned_faculty = faculty_list[faculty_index % num_faculty]
                if assigned_faculty not in faculty_assignments:
                    faculty_assignments[assigned_faculty] = []
                faculty_assignments[assigned_faculty].append((exam_date, slot_idx))
                faculty_index += 1
        
        # Create one row per faculty with all assignment slots
        data = []
        for sr, faculty_name in enumerate(faculty_list, 1):
            row = {
                "Sr. No.": sr,
                "Supervisor Name": faculty_name
            }
            
            # Add columns for each date-slot combination
            for exam_date in date_range:
                for slot_idx in range(len(time_slots)):
                    is_morning = slot_idx < morning_blocks
                    period = "M" if is_morning else "E"
                    slot_time = time_slots[slot_idx] if slot_idx < len(time_slots) else ""
                    
                    # Create column name with date, M/E, and slot timing
                    col_name = f"{exam_date.strftime('%d-%m-%Y')}_{period}_{slot_time}"
                    
                    # Check assignment
                    if (exam_date, slot_idx) in faculty_assignments.get(faculty_name, []):
                        row[col_name] = "✓"
                    else:
                        row[col_name] = ""
            
            data.append(row)
        
        # Create dataframe with wide format (one row per faculty)
        df = pd.DataFrame(data)

        st.subheader(f"Supervision Chart ({exam_start_date.strftime('%d-%m-%Y')} to {exam_end_date.strftime('%d-%m-%Y')})")
        total_exam_days = (exam_end_date - exam_start_date).days + 1
        st.write(f"**Total Days in Range:** {total_exam_days} | **Exam Days (after excluding {len(holidays)} holidays):** {len(date_range)} | **Time Slots/Day:** {len(time_slots)} | **Total Supervision Slots:** {total_slots}")
        st.write(f"**Total Faculty:** {num_faculty} | **Faculty Allocation:** Rotational (each faculty assigned to {total_slots // num_faculty + (1 if total_slots % num_faculty != 0 else 0)} slots)")
        st.dataframe(df, use_container_width=True)

        # Prepare data for exports with custom headers
        # Build multi-row header structure
        header_row1 = ["Sr. No.", "Supervisor Name"]
        header_row2 = ["", ""]
        header_row3 = ["", ""]
        
        for exam_date in date_range:
            date_str = exam_date.strftime("%d-%m-%Y")
            # Morning columns
            for slot_idx in range(min(morning_blocks, len(time_slots))):
                header_row1.append(date_str)
                header_row2.append("M")
                header_row3.append(time_slots[slot_idx])
            # Evening columns
            for slot_idx in range(morning_blocks, len(time_slots)):
                header_row1.append(date_str)
                header_row2.append("E")
                header_row3.append(time_slots[slot_idx])
        
        # Build CSV with multi-row headers
        csv_buffer = io.StringIO()
        csv_buffer.write("VVP Institute of engineering and Technology, Solapur\n")
        csv_buffer.write("Internal Supervision Chart\n")
        csv_buffer.write(f"Period: {exam_start_date.strftime('%d-%m-%Y')} to {exam_end_date.strftime('%d-%m-%Y')}\n")
        csv_buffer.write(f"Morning Blocks: {morning_blocks} | Evening Blocks: {evening_blocks}\n\n")
        
        # Write multi-row headers
        csv_buffer.write(",".join(str(h) for h in header_row1) + "\n")
        csv_buffer.write(",".join(str(h) for h in header_row2) + "\n")
        csv_buffer.write(",".join(str(h) for h in header_row3) + "\n")
        
        # Write data rows
        for idx, row in df.iterrows():
            row_values = [str(row["Sr. No."]), row["Supervisor Name"]]
            for col in df.columns[2:]:  # Skip Sr. No. and Supervisor Name
                row_values.append(row[col])
            csv_buffer.write(",".join(row_values) + "\n")
        
        csv_content = csv_buffer.getvalue()
        
        st.download_button(
            label="Download CSV",
            data=csv_content,
            file_name="supervision_chart.csv",
            mime="text/csv"
        )

        # Word Download with multi-row headers
        doc = Document()
        doc.add_heading("VVP Institute of engineering and Technology, Solapur", 0)
        doc.add_heading("Internal Supervision Chart", level=1)
        doc.add_paragraph(f"Period: {exam_start_date.strftime('%d-%m-%Y')} to {exam_end_date.strftime('%d-%m-%Y')}")
        doc.add_paragraph(f"Morning Blocks: {morning_blocks} | Evening Blocks: {evening_blocks}")
        doc.add_paragraph("")

        num_cols = len(df.columns)
        table = doc.add_table(rows=len(df)+3, cols=num_cols)
        table.style = 'Light Grid Accent 1'

        # Add header row 1 (Dates and merged Sr.No./Supervisor Name)
        for j, header_val in enumerate(header_row1):
            table.rows[0].cells[j].text = header_val
        
        # Add header row 2 (M/E)
        for j, header_val in enumerate(header_row2):
            table.rows[1].cells[j].text = header_val
        
        # Add header row 3 (Time slots)
        for j, header_val in enumerate(header_row3):
            table.rows[2].cells[j].text = header_val

        # Add data rows
        for i in range(len(df)):
            for j in range(num_cols):
                if j < 2:
                    table.rows[i+3].cells[j].text = str(df.iloc[i, j])
                else:
                    col_name = df.columns[j]
                    table.rows[i+3].cells[j].text = str(df.iloc[i, j])

        # Save Word to bytes buffer
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        st.download_button(
            label="Download Word",
            data=doc_buffer.getvalue(),
            file_name="supervision_chart.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # PDF Download with multi-row headers using reportlab
        pdf_buffer = io.BytesIO()
        doc = SimpleDocTemplate(pdf_buffer, pagesize=A4, topMargin=0.5*inch, bottomMargin=0.5*inch, leftMargin=0.3*inch, rightMargin=0.3*inch)
        story = []
        
        # Styles
        style = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=style['Heading1'],
            fontSize=14,
            textColor='#1e3c72',
            spaceAfter=6,
            alignment=1  # CENTER
        )
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=style['Heading2'],
            fontSize=12,
            textColor='#2a5298',
            spaceAfter=4,
            alignment=1  # CENTER
        )
        
        # Add titles
        story.append(Paragraph("VVP Institute of Engineering and Technology, Solapur", title_style))
        story.append(Paragraph("Internal Supervision Chart", subtitle_style))
        story.append(Paragraph(f"Period: {exam_start_date.strftime('%d-%m-%Y')} to {exam_end_date.strftime('%d-%m-%Y')}", style['Normal']))
        story.append(Paragraph(f"Morning Blocks: {morning_blocks} | Evening Blocks: {evening_blocks}", style['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Create table data
        df_pdf = df.copy()
        # Replace tick marks with 'X' 
        for col in df_pdf.columns:
            df_pdf[col] = df_pdf[col].astype(str).str.replace('✓', 'X')
        
        # Build table data with headers
        table_data = [header_row1, header_row2, header_row3]
        for idx, row in df_pdf.iterrows():
            row_data = [str(row["Sr. No."]), row["Supervisor Name"]]
            for col in df_pdf.columns[2:]:
                row_data.append(str(row[col])[:10])
            table_data.append(row_data)
        
        # Create table
        table = Table(table_data, colWidths=[0.6*inch if i < 2 else 0.35*inch for i in range(len(df_pdf.columns))])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 2), '#1e3c72'),
            ('TEXTCOLOR', (0, 0), (-1, 2), 'white'),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 2), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 2), 8),
            ('BOTTOMPADDING', (0, 0), (-1, 2), 4),
            ('GRID', (0, 0), (-1, -1), 0.5, 'black'),
            ('FONTSIZE', (0, 3), (-1, -1), 7),
            ('ROWBACKGROUNDS', (0, 3), (-1, -1), ['white', '#f0f0f0']),
        ]))
        story.append(table)
        
        # Build PDF
        doc.build(story)
        pdf_buffer.seek(0)
        
        st.download_button(
            label="Download PDF",
            data=pdf_buffer.getvalue(),
            file_name="supervision_chart.pdf",
            mime="application/pdf"
        )

# Footer
st.markdown("---")
st.markdown("""
    <div class='footer-container'>
        <div class='footer-text'>Prepared by</div>
        <div style='font-size: 16px; font-weight: bold;'>Prof. Amir M. Usman Wagdarikar</div>
        <div style='font-size: 13px;'>Asst. Prof., Department of Electronics and Telecommunication</div>
        <div style='margin-top: 10px; font-size: 12px; opacity: 0.9;'>VVP Institute of Engineering and Technology, Solapur</div>
    </div>
""", unsafe_allow_html=True)